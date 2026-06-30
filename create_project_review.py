import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', '.cache', 'codex-runtimes', 'codex-primary-runtime', 'dependencies', 'python'))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Design tokens (compact_reference_guide preset) ───
FONT_BASE = "Calibri"
FONT_SIZE_BASE = Pt(11)
COLOR_H1 = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2 = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H3 = RGBColor(0x1F, 0x4D, 0x78)
COLOR_INK = RGBColor(0x0B, 0x25, 0x45)
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)
COLOR_TABLE_HEADER = RGBColor(0xE8, 0xEE, 0xF5)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PAGE_WIDTH_IN = 6.5
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = 120


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tcPr.append(shading)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    """Set cell margins in DXA."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professional table with header styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table geometry
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(TABLE_INDENT_DXA))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.font.name = FONT_BASE
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = COLOR_INK
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.name = FONT_BASE
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_INK
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)
    
    # Set borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D9D9D9')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    return table


def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = FONT_BASE
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = COLOR_H1
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_H2
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_H3
    return p


def add_paragraph_styled(doc, text, bold=False, size=None, color=None, alignment=None, space_after=Pt(6)):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = FONT_BASE
    run.font.size = size or FONT_SIZE_BASE
    run.font.color.rgb = color or COLOR_INK
    run.font.bold = bold
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_BASE
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_INK
    p.paragraph_format.space_after = Pt(4)
    return p


def main():
    doc = Document()
    
    # ─── Page setup ───
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    
    # ─── Set default font ───
    style = doc.styles['Normal']
    style.font.name = FONT_BASE
    style.font.size = FONT_SIZE_BASE
    style.paragraph_format.space_after = Pt(6)
    
    # ═══════════════════════════════════════════════════════════════
    # COVER / TITLE PAGE
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph()  # spacing
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LiveIn (HongShu)")
    run.font.name = FONT_BASE
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_H1
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("项目深度审查报告")
    run.font.name = FONT_BASE
    run.font.size = Pt(20)
    run.font.color.rgb = COLOR_INK
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 SpringBoot + Vue + Uniapp + AI 大模型 的仿小红书全栈项目")
    run.font.name = FONT_BASE
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_MUTED
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"审查日期: {datetime.date.today().strftime('%Y-%m-%d')}")
    run.font.name = FONT_BASE
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_MUTED
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "目录", level=1)
    toc_items = [
        "一、项目概述",
        "二、后端技术栈",
        "三、前端技术栈",
        "四、其他技术",
        "五、项目特色",
        "六、数据库设计",
        "七、API 接口文档",
        "八、项目配置要点",
        "九、技术架构图",
    ]
    for item in toc_items:
        add_paragraph_styled(doc, item, size=Pt(11), color=COLOR_INK, space_after=Pt(4))
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 一、项目概述
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "一、项目概述", level=1)
    
    add_styled_table(doc,
        ["属性", "说明"],
        [
            ["项目名称", "来因LiveIn (HongShu)"],
            ["项目定位", "基于 SpringBoot + Vue + Uniapp + AI大模型 的1:1高仿小红书全栈项目"],
            ["开源协议", "MIT"],
            ["Gitee", "https://gitee.com/Maverick_Ma/hongshu"],
            ["GitHub", "https://github.com/Ma-YongJian/HongShu"],
        ],
        col_widths=[1.5, 5.0]
    )
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "项目包含三端完整实现:", bold=True, space_after=Pt(4))
    add_bullet(doc, "hongshu-master — 后端服务 (Java)")
    add_bullet(doc, "hongshu-ui-master/hongshu-web — Web端前台 (Vue3 + TypeScript)")
    add_bullet(doc, "hongshu-ui-master/hongshu-admin — 管理后台 (Vue3 + Element Plus, 基于若依框架)")
    add_bullet(doc, "hongshu-ui-master/hongshu-app — 移动端 (UniApp, 支持H5/APP/微信小程序)")
    
    # ═══════════════════════════════════════════════════════════════
    # 二、后端技术栈
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "二、后端技术栈", level=1)
    
    add_heading_styled(doc, "2.1 核心框架", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["Java", "1.8", "开发语言"],
            ["Spring Boot", "2.5.15", "应用框架"],
            ["Spring Framework", "5.3.33", "Spring核心"],
            ["Spring Security", "随SpringBoot", "权限认证, @PreAuthorize 注解鉴权"],
            ["MyBatis-Plus", "3.5.2", "ORM框架, 简化数据库操作"],
            ["Swagger (springfox)", "3.0.0", "API文档自动生成"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "2.2 数据层", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["MySQL/MariaDB", "-", "主数据库 (驱动: mariadb-java-client 3.1.4)"],
            ["Druid", "1.2.20", "阿里巴巴数据库连接池, 支持慢SQL监控"],
            ["Redis", "7.x", "缓存、Token存储、会话管理 (Lettuce客户端)"],
            ["ElasticSearch", "8.1.0", "全文搜索引擎, 笔记搜索与推荐"],
            ["PageHelper", "1.4.7", "MyBatis物理分页插件"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "2.3 安全与认证", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["JJWT", "0.9.1", "JWT Token生成与验证, 双Token无感刷新机制"],
            ["Kaptcha", "2.3.3", "验证码生成 (支持数学计算/字符两种模式)"],
            ["JSEncrypt", "3.3.2", "前端RSA加密传输密码"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "2.4 工具库", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["Hutool", "5.7.22", "Java工具集 (JSON处理、加密、日期等)"],
            ["Fastjson2", "2.0.43", "阿里JSON解析器"],
            ["Gson", "2.9.0", "Google JSON解析器"],
            ["Guava", "20.0", "Google核心库 (缓存、集合等)"],
            ["Lombok", "1.18.24", "简化Java代码 (Getter/Setter/Builder等)"],
            ["Apache POI", "4.1.2", "Excel导入导出"],
            ["Apache Velocity", "2.3", "代码生成模板引擎"],
            ["Commons IO", "2.13.0", "IO工具类"],
            ["OSHI", "6.5.0", "系统信息采集 (CPU/内存/磁盘监控)"],
            ["UserAgentUtils", "1.21", "客户端浏览器/操作系统解析"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "2.5 实时通信", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["WebSocket (javax.websocket)", "-", "即时通讯, @ServerEndpoint 注解实现"],
            ["Spring Boot WebSocket Starter", "3.1.7", "WebSocket服务端支持"],
        ],
        col_widths=[2.5, 1.0, 3.0]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "2.6 文件存储", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["七牛云 SDK", "7.7.0", "七牛云对象存储"],
            ["MinIO", "-", "本地私有对象存储"],
            ["阿里云/腾讯云", "-", "云端对象存储"],
            ["本地存储", "-", "本地文件系统存储"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )
    add_paragraph_styled(doc, "存储策略支持动态切换, 通过后台 oss.type 配置项控制 (0:本地, 1:七牛云, 2:MinIO, 3:阿里云, 4:腾讯云)", 
                          size=Pt(10), color=COLOR_MUTED, space_after=Pt(8))
    
    doc.add_paragraph()
    add_heading_styled(doc, "2.7 后端模块结构", level=2)
    code_text = (
        "hongshu-master/\n"
        "├── hongshu-common      # 公共工具模块 (注解、常量、工具类、校验器)\n"
        "├── hongshu-framework   # 框架核心 (Spring配置、AOP、拦截器、安全配置)\n"
        "├── hongshu-system      # 系统模块 (系统管理相关Entity/Mapper/Service)\n"
        "├── hongshu-web         # 业务模块 (核心业务Entity/Mapper/Service/WebSocket)\n"
        "├── hongshu-quartz      # 定时任务模块 (Quartz调度、任务日志)\n"
        "└── hongshu-server      # 启动模块 (所有Controller、配置文件、主启动类)"
    )
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLOR_INK
    p.paragraph_format.space_after = Pt(8)
    
    # ═══════════════════════════════════════════════════════════════
    # 三、前端技术栈
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "三、前端技术栈", level=1)
    
    add_heading_styled(doc, "3.1 Web端前台 (hongshu-web)", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["Vue", "3.3.4", "前端框架"],
            ["TypeScript", "5.0.2", "类型安全的JavaScript"],
            ["Vite", "4.4.5", "构建工具"],
            ["Element Plus", "2.4.1+", "UI组件库"],
            ["Pinia", "2.1.7", "状态管理"],
            ["Pinia PersistedState", "3.2.0", "状态持久化"],
            ["Axios", "1.5.1", "HTTP请求库"],
            ["Vue Router", "4.2.5", "路由管理"],
            ["VueUse", "12.3.0+", "Vue组合式工具集"],
            ["vue-waterfall-plugin-next", "2.2.4", "瀑布流布局"],
            ["vue3-lazy", "1.0.0-alpha.1", "图片懒加载"],
            ["vue-picture-cropper", "0.7.0", "图片裁剪"],
            ["GoEasy", "2.10.14", "IM消息服务SDK"],
            ["zego-zim-web", "2.12.0", "即时通讯SDK"],
            ["Less / Sass", "4.2.0 / 1.77.1", "CSS预处理器"],
        ],
        col_widths=[2.2, 1.3, 3.0]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "3.2 管理后台 (hongshu-admin)", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["Vue", "3.3.9", "前端框架"],
            ["Vite", "5.0.4", "构建工具"],
            ["Element Plus", "2.7.5+", "UI组件库 (含 @element-plus/icons-vue)"],
            ["Pinia", "2.1.7", "状态管理"],
            ["Vuex", "4.0.2", "部分模块仍使用Vuex"],
            ["Axios", "0.27.2", "HTTP请求库"],
            ["ECharts", "5.4.3", "数据可视化图表"],
            ["Vue-Quill", "1.2.0", "富文本编辑器"],
            ["Vue-Cropper", "1.1.1", "图片裁剪组件"],
            ["JSEncrypt", "3.3.2", "RSA加密"],
            ["SortableJS", "1.15.2", "拖拽排序"],
            ["Fuse.js", "6.6.2", "模糊搜索"],
            ["NProgress", "0.2.0", "进度条"],
            ["js-cookie", "3.0.5", "Cookie操作"],
            ["file-saver", "2.0.5", "文件下载"],
        ],
        col_widths=[2.0, 1.3, 3.2]
    )
    add_paragraph_styled(doc, "管理后台基于若依 (RuoYi-Vue) 框架搭建", size=Pt(10), color=COLOR_MUTED)
    
    doc.add_paragraph()
    add_heading_styled(doc, "3.3 移动端 (hongshu-app)", level=2)
    add_styled_table(doc,
        ["技术", "版本", "用途"],
        [
            ["UniApp", "-", "跨平台移动框架 (H5/APP/微信小程序)"],
            ["Vue", "2.x", "UniApp默认Vue版本"],
            ["ThorUI", "-", "移动端UI组件库"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "3.4 前端页面结构", level=2)
    add_paragraph_styled(doc, "Web端页面:", bold=True, space_after=Pt(2))
    add_paragraph_styled(doc, "首页(dashboard)、动态关注、搜索、发布笔记、消息中心(评论/点赞收藏/关注/私信)、用户主页、专辑管理", space_after=Pt(6))
    
    add_paragraph_styled(doc, "管理后台页面:", bold=True, space_after=Pt(2))
    add_paragraph_styled(doc, "数据仪表盘、会员管理、笔记管理(含审核)、评论管理、专辑管理、标签管理、导航栏管理、图片管理、爬虫管理、系统配置、日志监控、缓存管理、服务器监控、字典管理、菜单管理、角色管理、部门管理、岗位管理", space_after=Pt(6))
    
    add_paragraph_styled(doc, "App端页面:", bold=True, space_after=Pt(2))
    add_paragraph_styled(doc, "首页、发现/兴趣、消息、个人中心、登录/注册、聊天、设置、专辑、搜索、发布、标签、群聊等30+页面", space_after=Pt(6))
    
    # ═══════════════════════════════════════════════════════════════
    # 四、其他技术
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "四、其他技术", level=1)
    add_styled_table(doc,
        ["技术", "用途"],
        [
            ["Docker + Docker Compose", "容器化部署"],
            ["Nginx", "反向代理、静态资源服务、负载均衡"],
            ["SpringBoot Admin", "应用监控"],
            ["爬虫 (Python)", "图片/内容爬取"],
            ["高德地图", "LBS定位、同城推荐"],
            ["支付宝/微信支付", "在线支付"],
            ["阿里云/腾讯云短信", "短信验证码"],
            ["百度千帆大模型", "图片/视频内容审核"],
        ],
        col_widths=[2.5, 4.0]
    )
    
    # ═══════════════════════════════════════════════════════════════
    # 五、项目特色
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "五、项目特色", level=1)
    
    features = [
        ("智能推荐算法", "三套算法 (轻量级推荐 + 协同过滤 + ES检索) 动态切换, 新用户冷启动优化, 推荐响应性能 < 100ms"),
        ("高并发优化", "Redis + RocketMQ 双通道处理, 批量落库策略, 性能提升20倍, 数据库压力降低90%"),
        ("AI大模型集成", "支持 ChatGPT + DeepSeek 等8+模型, 智能创作、AI对话系统"),
        ("安全认证", "双Token无感刷新 (accessToken + refreshToken), RSA加密传输密码"),
        ("多云存储", "本地 + MinIO + 七牛云 + 阿里云 + 腾讯云, 后台动态切换存储策略"),
        ("智能搜索", "ElasticSearch + MySQL 双引擎, 多字段权重排序、模糊匹配"),
        ("即时通讯", "WebSocket实时消息, 离线消息存储、多端同步"),
        ("社交电商闭环", "笔记种草 → 商品转化 → 在线支付, 价格区间推荐、同城优先"),
        ("多端适配", "Web + APP + 微信小程序, 代码复用, 统一API"),
        ("自定义注解体系", "@Anonymous、@NoLoginIntercept、@RateLimiter、@RepeatSubmit、@Sensitive、@Log、@DataScope、@AuthorityVerify 等"),
    ]
    
    add_styled_table(doc,
        ["特色模块", "技术实现与创新点"],
        [[f, d] for f, d in features],
        col_widths=[1.8, 4.7]
    )
    
    # ═══════════════════════════════════════════════════════════════
    # 六、数据库设计
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "六、数据库设计", level=1)
    add_paragraph_styled(doc, "共37张表, 分为系统管理表和业务表两大类:", space_after=Pt(8))
    
    add_heading_styled(doc, "6.1 系统管理表 (18张)", level=2)
    sys_tables = [
        ["sys_config", "系统配置"],
        ["sys_date", "日期"],
        ["sys_dept", "部门"],
        ["sys_dict_data", "字典数据"],
        ["sys_dict_type", "字典类型"],
        ["sys_job", "定时任务"],
        ["sys_job_log", "任务日志"],
        ["sys_logininfor", "登录日志"],
        ["sys_menu", "菜单权限"],
        ["sys_notice", "通知公告"],
        ["sys_oper_log", "操作日志"],
        ["sys_post", "岗位"],
        ["sys_role", "角色"],
        ["sys_role_dept", "角色部门关联"],
        ["sys_role_menu", "角色菜单关联"],
        ["sys_user", "系统用户"],
        ["sys_user_post", "用户岗位关联"],
        ["sys_user_role", "用户角色关联"],
    ]
    add_styled_table(doc, ["表名", "说明"], sys_tables, col_widths=[2.5, 4.0])
    
    doc.add_paragraph()
    add_heading_styled(doc, "6.2 业务表 (19张)", level=2)
    biz_tables = [
        ["web_user", "前台用户"],
        ["web_note", "笔记"],
        ["web_picture", "图片"],
        ["web_picture_sort", "图片分类"],
        ["web_comment", "评论"],
        ["web_comment_sync", "评论同步"],
        ["web_like_or_collect", "点赞收藏"],
        ["web_follow", "关注"],
        ["web_chat", "聊天消息"],
        ["web_chat_user_relation", "聊天用户关联"],
        ["web_tag", "标签"],
        ["web_tag_note_relation", "标签笔记关联"],
        ["web_album", "专辑"],
        ["web_album_note_relation", "专辑笔记关联"],
        ["web_navbar", "导航栏/分类"],
        ["web_login_log", "前台登录日志"],
        ["web_oper_log", "前台操作日志"],
        ["web_user_note_relation", "用户笔记关联"],
        ["web_visit", "访问记录"],
    ]
    add_styled_table(doc, ["表名", "说明"], biz_tables, col_widths=[2.5, 4.0])
    
    # ═══════════════════════════════════════════════════════════════
    # 七、API接口文档
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "七、API 接口文档", level=1)
    
    # 7.1 认证模块
    add_heading_styled(doc, "7.1 认证模块 /web/auth", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["POST", "/web/auth/login", "用户登录", "否"],
            ["POST", "/web/auth/loginByCode", "验证码登录", "否"],
            ["POST", "/web/auth/register", "用户注册", "否"],
            ["POST", "/web/auth/isRegister", "检查用户是否已注册", "是"],
            ["GET", "/web/auth/getUserInfoByToken", "根据Token获取用户信息", "是"],
            ["GET", "/web/auth/loginOut", "退出登录", "是"],
            ["POST", "/web/auth/updatePassword", "修改密码", "是"],
            ["GET", "/web/auth/refreshToken", "刷新Token", "否"],
        ],
        col_widths=[0.7, 2.5, 2.0, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.2 笔记模块 /web/note", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["GET", "/web/note/getNoteById?noteId=", "获取笔记详情", "否"],
            ["POST", "/web/note/saveNoteByDTO", "新增笔记(含图片上传)", "是"],
            ["POST", "/web/note/updateNoteByDTO", "更新笔记(含图片上传)", "是"],
            ["POST", "/web/note/deleteNoteByIds", "批量删除笔记", "是"],
            ["GET", "/web/note/getHotPage/{page}/{size}", "获取热门笔记分页", "是"],
            ["GET", "/web/note/pinnedNote?noteId=", "置顶笔记", "是"],
        ],
        col_widths=[0.7, 2.5, 2.0, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.3 评论模块 /web/comment", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["GET", "/web/comment/getOneCommentByNoteId/{page}/{size}", "获取笔记的一级评论", "是"],
            ["GET", "/web/comment/getCommentById?commentId=", "获取单个评论详情", "是"],
            ["POST", "/web/comment/saveCommentByDTO", "发表评论", "是"],
            ["POST", "/web/comment/syncCommentByIds", "同步评论(批量)", "是"],
            ["GET", "/web/comment/getTwoCommentByOneCommentId/{page}/{size}", "获取二级评论", "是"],
            ["GET", "/web/comment/getNoticeComment/{page}/{size}", "获取通知评论", "是"],
            ["GET", "/web/comment/getCommentWithCommentByNoteId/{page}/{size}", "获取带子评论的评论树", "是"],
            ["GET", "/web/comment/scrollComment?commentId=", "自动滚动定位到评论", "是"],
            ["GET", "/web/comment/deleteCommentById?commentId=", "删除评论", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.4 关注模块 /web/follower", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["GET", "/web/follower/getFollowTrend/{page}/{size}", "获取关注用户的动态", "是"],
            ["GET", "/web/follower/getFriend/{page}/{size}?uid=&type=", "获取关注/粉丝列表", "是"],
            ["GET", "/web/follower/followById?followerId=", "关注用户", "是"],
            ["GET", "/web/follower/isFollow?followerId=", "检查是否已关注", "是"],
            ["GET", "/web/follower/getNoticeFollower/{page}/{size}", "获取最新关注通知", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.5 点赞/收藏模块 /web/likeOrCollection", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["POST", "/web/likeOrCollection/likeOrCollectionByDTO", "点赞或收藏", "是"],
            ["POST", "/web/likeOrCollection/isLikeOrCollection", "检查是否已点赞/收藏", "是"],
            ["GET", "/web/likeOrCollection/getNoticeLikeOrCollection/{page}/{size}", "获取通知", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.6 即时通讯模块 /web/im/chat", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["POST", "/web/im/chat/sendMsg", "发送消息", "否"],
            ["GET", "/web/im/chat/getAllChatRecord/{page}/{size}?acceptUid=", "获取聊天记录", "是"],
            ["GET", "/web/im/chat/getChatUserList", "获取聊天用户列表", "是"],
            ["GET", "/web/im/chat/getCountMessage", "获取未读消息统计", "是"],
            ["GET", "/web/im/chat/clearMessageCount?sendUid=&type=", "清除消息计数", "是"],
            ["GET", "/web/im/chat/closeChat/{sendUid}", "关闭聊天", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.7 用户模块 /web/user", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["GET", "/web/user/getUserById?userId=", "获取用户信息", "是"],
            ["POST", "/web/user/updateUser", "更新用户信息", "是"],
            ["GET", "/web/user/getTrendByUser/{page}/{size}?userId=&type=", "获取用户动态", "是"],
            ["GET", "/web/user/getUserByKeyword/{page}/{size}?keyword=", "搜索用户", "是"],
            ["GET", "/web/user/saveUserSearchRecord?keyword=", "保存搜索记录", "是"],
            ["GET", "/web/user/getUserList", "会员列表(管理端)", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.8 专辑模块 /web/album", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["GET", "/web/album/getAlbumByUserId/{page}/{size}?userId=", "获取用户专辑", "是"],
            ["POST", "/web/album/saveAlbum", "创建专辑", "是"],
            ["GET", "/web/album/getAlbumById?albumId=", "获取专辑详情", "是"],
            ["GET", "/web/album/deleteAlbum?albumId=", "删除专辑", "是"],
            ["POST", "/web/album/updateAlbum", "更新专辑", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.9 标签模块 /web/tag", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["GET", "/web/tag/getHotTagList", "获取热门标签", "是"],
            ["GET", "/web/tag/getTagByKeyword/{page}/{size}?keyword=", "搜索标签", "是"],
            ["GET", "/web/tag/getTagById?tagId=", "获取标签详情", "是"],
            ["GET", "/web/tag/getNoteByTagId/{page}/{size}?tagId=&type=", "获取标签下的笔记", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.10 分类模块 /web/category", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["GET", "/web/category/getCategoryTreeData", "获取树形分类数据", "否"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.11 ES搜索模块 /web/es/note", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["POST", "/web/es/note/getNoteByDTO/{page}/{size}", "ES全文搜索笔记", "否"],
            ["POST", "/web/es/note/getCategoryAgg", "ES分类聚合查询", "否"],
            ["GET", "/web/es/note/getRecommendNote/{page}/{size}", "ES推荐笔记", "否"],
            ["GET", "/web/es/note/getHotNote/{page}/{size}", "ES热门笔记", "是"],
            ["POST", "/web/es/note/addNote", "ES新增索引", "是"],
            ["POST", "/web/es/note/updateNote", "ES更新索引", "是"],
            ["GET", "/web/es/note/deleteNote/{noteId}", "ES删除索引", "是"],
            ["POST", "/web/es/note/addNoteBulkData", "ES批量导入", "否"],
            ["POST", "/web/es/note/refreshNoteData", "ES重置数据", "否"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.12 ES搜索记录 /web/es/record", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["GET", "/web/es/record/getRecordByKeyWord", "获取搜索记录", "是"],
            ["GET", "/web/es/record/getHotRecord", "获取热门搜索", "否"],
            ["POST", "/web/es/record/addRecord", "添加搜索记录", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.13 文件上传模块 /web/oss", level=2)
    add_styled_table(doc,
        ["方法", "路径", "说明", "需登录"],
        [
            ["POST", "/web/oss/save/{type}", "单文件上传", "是"],
            ["POST", "/web/oss/saveBatch", "批量文件上传", "是"],
            ["GET", "/web/oss/delete?path=", "删除文件", "是"],
            ["POST", "/web/oss/deleteBatch", "批量删除文件", "是"],
        ],
        col_widths=[0.7, 3.0, 1.5, 0.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.14 WebSocket", level=2)
    add_styled_table(doc,
        ["路径", "说明"],
        [
            ["/web/ws/{uid}", "WebSocket连接端点, uid为用户ID"],
        ],
        col_widths=[2.5, 4.0]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "7.15 管理后台接口", level=2)
    
    add_heading_styled(doc, "系统管理", level=3)
    add_styled_table(doc,
        ["方法", "路径", "说明", "权限"],
        [
            ["POST", "/login", "管理员登录", "公开"],
            ["POST", "/register", "管理员注册", "公开"],
            ["GET", "/getInfo", "获取当前管理员信息", "需认证"],
            ["GET", "/user/list", "用户列表", "system:user:list"],
            ["POST", "/user", "新增用户", "system:user:add"],
            ["PUT", "/user", "修改用户", "system:user:edit"],
            ["DELETE", "/user/{ids}", "删除用户", "system:user:remove"],
            ["GET", "/role/list", "角色列表", "system:role:list"],
            ["GET", "/menu/list", "菜单列表", "system:menu:list"],
            ["GET", "/dept/list", "部门列表", "system:dept:list"],
            ["GET", "/post/list", "岗位列表", "system:post:list"],
            ["GET", "/dict/type", "字典类型", "system:dict:list"],
            ["GET", "/config/list", "参数配置", "system:config:list"],
            ["GET", "/notice/list", "通知公告", "system:notice:list"],
        ],
        col_widths=[0.7, 2.0, 2.0, 1.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "业务管理", level=3)
    add_styled_table(doc,
        ["方法", "路径", "说明", "权限"],
        [
            ["GET", "/note/list", "笔记列表", "web:note:list"],
            ["GET", "/note/unAuditList", "未审核笔记", "web:note:list"],
            ["PUT", "/note/auditNote", "审核笔记(通过/拒绝)", "-"],
            ["POST", "/note", "新增笔记", "web:note:add"],
            ["PUT", "/note", "修改笔记", "web:note:edit"],
            ["DELETE", "/note/{ids}", "删除笔记", "web:note:remove"],
            ["POST", "/note/export", "导出笔记", "web:note:export"],
            ["GET", "/member/list", "会员列表", "web:member:list"],
            ["POST", "/member", "新增会员", "web:member:add"],
            ["PUT", "/member", "修改会员", "web:member:edit"],
            ["DELETE", "/member/{ids}", "删除会员", "web:member:remove"],
            ["POST", "/member/export", "导出会员", "web:member:export"],
        ],
        col_widths=[0.7, 2.0, 2.0, 1.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "数据统计", level=3)
    add_styled_table(doc,
        ["方法", "路径", "说明"],
        [
            ["GET", "/index/init", "首页初始化(访问量/用户数/笔记数/评论数)"],
            ["GET", "/index/getVisitByWeek", "近一周访问量趋势"],
            ["GET", "/index/getBlogCountByTag", "每个标签下文章数量"],
            ["GET", "/index/getBlogContributeCount", "一年内文章贡献度"],
            ["GET", "/statistics/index/total", "总计数据"],
            ["GET", "/statistics/index/line", "折线图数据"],
            ["GET", "/statistics/index/raddar", "雷达图数据"],
            ["GET", "/statistics/index/pie", "饼图数据"],
            ["GET", "/statistics/index/bar", "柱状图数据"],
        ],
        col_widths=[0.7, 3.0, 2.8]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "系统运维", level=3)
    add_styled_table(doc,
        ["方法", "路径", "说明"],
        [
            ["GET", "/monitor/cache", "缓存信息"],
            ["GET", "/monitor/server", "服务器信息(CPU/内存/磁盘)"],
            ["GET", "/monitor/logininfor/list", "登录日志"],
            ["GET", "/monitor/operlog/list", "操作日志"],
            ["GET", "/monitor/online", "在线用户"],
            ["GET", "/monitor/job/list", "定时任务"],
            ["POST", "/systemConfig/getSystemConfig", "获取系统配置"],
            ["POST", "/systemConfig/editSystemConfig", "修改系统配置"],
            ["POST", "/systemConfig/cleanRedisByKey", "按Key清理Redis缓存"],
        ],
        col_widths=[0.7, 3.0, 2.8]
    )
    
    # ═══════════════════════════════════════════════════════════════
    # 八、项目配置要点
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "八、项目配置要点", level=1)
    add_styled_table(doc,
        ["配置项", "默认值", "说明"],
        [
            ["服务端口", "8080", "Spring Boot HTTP端口"],
            ["前端端口", "80", "Vite开发服务器端口"],
            ["数据库", "MariaDB 3308/hongshu_springboot", "主数据库"],
            ["Redis", "localhost:6379", "缓存服务"],
            ["ES", "localhost:9200", "搜索引擎"],
            ["Token有效期", "30分钟", "JWT Token过期时间"],
            ["文件上传限制", "单文件10MB, 总20MB", "上传大小限制"],
            ["文件存储路径", "I:/hongshu-ma/upload", "本地存储路径 (Windows)"],
            ["密码最大错误次数", "5次", "锁定前允许的错误次数"],
            ["密码锁定时间", "10分钟", "锁定时长"],
            ["XSS过滤", "开启", "对 /system/*, /monitor/*, /tool/* 生效"],
        ],
        col_widths=[2.0, 2.5, 2.0]
    )
    
    # ═══════════════════════════════════════════════════════════════
    # 九、技术架构图
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "九、技术架构图", level=1)
    
    # Create a visual architecture table
    arch_data = [
        ["客户端层", "Web前台 (Vue3+TS)", "Admin管理后台 (Vue3+Element)", "UniApp (H5/APP/小程序)"],
        ["", "↓", "↓", "↓"],
        ["代理层", "Nginx 反向代理", "", ""],
        ["", "↓", "", ""],
        ["后端服务层", "Spring Boot 2.5.15", "", ""],
        ["", "认证模块 (JWT) | 业务Controller | 系统管理 | 定时任务 (Quartz)", "", ""],
        ["", "Service / DAO (MyBatis-Plus)", "", ""],
        ["", "↓", "", ""],
        ["数据存储层", "MySQL/MariaDB", "Redis 7.x", "ES 8.x"],
        ["", "MinIO/七牛云", "RocketMQ", ""],
    ]
    
    arch_table = add_styled_table(doc,
        ["层级", "组件A", "组件B", "组件C"],
        arch_data,
        col_widths=[1.3, 2.2, 2.0, 1.0]
    )
    
    # Merge cells for architecture visualization
    # Row 5 (Spring Boot) - merge cols B, C, D
    row5 = arch_table.rows[5]
    row5.cells[1].merge(row5.cells[3])
    # Row 6 (Service/DAO) - merge cols B, C, D
    row6 = arch_table.rows[6]
    row6.cells[1].merge(row6.cells[3])
    # Row 2 (Nginx) - merge cols B, C, D
    row2 = arch_table.rows[2]
    row2.cells[1].merge(row2.cells[3])
    # Row 3 (arrow) - merge cols B, C, D
    row3 = arch_table.rows[3]
    row3.cells[1].merge(row3.cells[3])
    # Row 7 (arrow) - merge cols B, C, D
    row7 = arch_table.rows[7]
    row7.cells[1].merge(row7.cells[3])
    
    # ─── Footer ───
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("来因LiveIn 项目审查报告")
    run.font.name = FONT_BASE
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED
    
    # ─── Save ───
    output_path = os.path.join(os.path.dirname(__file__), "来因LiveIn_项目深度审查报告.docx")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    main()
